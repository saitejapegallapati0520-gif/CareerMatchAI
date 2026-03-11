"""Resume tailoring and DOCX generation module."""

from __future__ import annotations

import json
import os
import re
import logging
from typing import Any, Dict, List

from docx import Document
from docx.shared import Pt
from dotenv import load_dotenv
from langchain_google_genai import ChatGoogleGenerativeAI


load_dotenv()
logger = logging.getLogger(__name__)


def _invoke_with_fallback(prompt: str) -> str:
    """Invoke Gemini with fallback models and return text content."""
    try:
        api_key = os.getenv("GEMINI_API_KEY", "").strip()
        if not api_key:
            raise ValueError("Missing GEMINI_API_KEY in environment.")
        model_candidates = ["gemini-1.5-flash", "gemini-2.0-flash", "gemini-flash-latest"]
        last_error: Exception | None = None
        for model_name in model_candidates:
            try:
                llm = ChatGoogleGenerativeAI(
                    model=model_name,
                    google_api_key=api_key,
                    temperature=0.1,
                )
                response = llm.invoke(prompt)
                return str(getattr(response, "content", "") or "")
            except Exception as model_exc:
                last_error = model_exc
                continue
        raise RuntimeError(str(last_error) if last_error else "Unknown LLM invocation error.")
    except Exception as exc:
        raise RuntimeError(f"Failed to invoke LLM client: {exc}") from exc


def _extract_json(text: str) -> Dict[str, Any]:
    """Extract first JSON object from model output."""
    try:
        cleaned = str(text or "").replace("```json", "").replace("```", "").strip()
        match = re.search(r"\{[\s\S]*\}", cleaned)
        if not match:
            return {}
        return json.loads(match.group(0))
    except Exception:
        return {}


def tailor_summary(summary: str, job_description: str) -> str:
    """Tailor professional summary to target job without fabricating facts."""
    try:
        logger.info("Tailoring summary")
        if not summary.strip():
            return ""
        prompt = f"""
Rewrite the resume summary to better align with the job description.
Strict rules:
1) Do not invent skills, titles, companies, or achievements.
2) Only rephrase, prioritize, and reorder details already present in the input summary.
3) Keep it ATS-friendly and concise (3-4 lines).

Return strict JSON: {{"tailored_summary": "..."}}.

INPUT SUMMARY:
{summary}

JOB DESCRIPTION:
{job_description}
"""
        content = _invoke_with_fallback(prompt)
        parsed = _extract_json(content)
        tailored = str(parsed.get("tailored_summary", "")).strip()
        return tailored if tailored else summary
    except Exception as exc:
        logger.warning("Error in tailor_summary: %s", exc)
        return summary


def tailor_experience(experience: List[str], job_description: str) -> List[str]:
    """Tailor experience bullets by reordering and rephrasing existing content."""
    try:
        logger.info("Tailoring experience")
        if not experience:
            return []

        experience_text = "\n".join([f"- {item}" for item in experience])
        prompt = f"""
Optimize these experience bullets for ATS alignment with the job description.
Strict rules:
1) Do not fabricate achievements or technologies not already present.
2) Keep each bullet truthful and concise.
3) Reorder bullets by relevance to the job.
4) Improve action verbs and clarity.

Return strict JSON:
{{"tailored_experience": ["bullet 1", "bullet 2", "..."]}}

EXPERIENCE BULLETS:
{experience_text}

JOB DESCRIPTION:
{job_description}
"""
        content = _invoke_with_fallback(prompt)
        parsed = _extract_json(content)
        tailored = parsed.get("tailored_experience", [])
        if isinstance(tailored, list):
            cleaned = [str(item).strip() for item in tailored if str(item).strip()]
            return cleaned if cleaned else experience
        return experience
    except Exception as exc:
        logger.warning("Error in tailor_experience: %s", exc)
        return experience


def tailor_skills(skills: List[str], job_description: str) -> Dict[str, List[str]]:
    """Tailor skills section by grouping matched, missing, and prioritized skills."""
    try:
        logger.info("Tailoring skills")
        normalized_skills = [str(s).strip() for s in (skills or []) if str(s).strip()]
        jd_text = str(job_description or "").lower()
        matched = [skill for skill in normalized_skills if skill.lower() in jd_text]
        missing = [skill for skill in normalized_skills if skill.lower() not in jd_text]
        prioritized = matched + [skill for skill in missing if skill not in matched]
        return {
            "matched_skills": matched,
            "missing_skills": missing,
            "prioritized_skills": prioritized,
        }
    except Exception as exc:
        logger.warning("Error in tailor_skills: %s", exc)
        return {"matched_skills": [], "missing_skills": [], "prioritized_skills": skills or []}


def tailor_projects(projects: List[str], job_description: str) -> List[str]:
    """Tailor project bullets by reordering and clarifying existing items only."""
    try:
        logger.info("Tailoring projects")
        if not projects:
            return []

        projects_text = "\n".join([f"- {item}" for item in projects])
        prompt = f"""
Reorder and rewrite project bullets to best match the target job.
Strict rules:
1) Do not add fabricated projects, skills, or outcomes.
2) Only use details present in the input.
3) Keep bullets concise and ATS-friendly.

Return strict JSON:
{{"tailored_projects": ["project bullet 1", "project bullet 2"]}}

PROJECT BULLETS:
{projects_text}

JOB DESCRIPTION:
{job_description}
"""
        content = _invoke_with_fallback(prompt)
        parsed = _extract_json(content)
        tailored = parsed.get("tailored_projects", [])
        if isinstance(tailored, list):
            cleaned = [str(item).strip() for item in tailored if str(item).strip()]
            return cleaned if cleaned else projects
        return projects
    except Exception as exc:
        logger.warning("Error in tailor_projects: %s", exc)
        return projects


def build_tailored_resume(resume_data: Dict[str, Any], job: Dict[str, Any]) -> Dict[str, Any]:
    """Build a complete tailored resume dictionary for a specific job."""
    try:
        logger.info("Building tailored resume object")
        structured = resume_data.get("structured_info", {}) if isinstance(resume_data.get("structured_info", {}), dict) else {}
        job_description = str(job.get("description", "")).strip()

        summary = str(structured.get("summary", "")).strip()
        skills = structured.get("skills", []) if isinstance(structured.get("skills", []), list) else []
        experience = structured.get("experience", []) if isinstance(structured.get("experience", []), list) else []
        projects = structured.get("projects", []) if isinstance(structured.get("projects", []), list) else []
        education = structured.get("education", []) if isinstance(structured.get("education", []), list) else []
        certifications = structured.get("certifications", []) if isinstance(structured.get("certifications", []), list) else []

        tailored_summary = tailor_summary(summary, job_description)
        tailored_experience = tailor_experience(experience, job_description)
        tailored_skill_map = tailor_skills(skills, job_description)
        tailored_projects = tailor_projects(projects, job_description)

        return {
            "header": {
                "name": str(structured.get("name", "")).strip(),
                "email": str(structured.get("email", "")).strip(),
                "target_role": str(job.get("title", "")).strip(),
            },
            "summary": tailored_summary,
            "skills": tailored_skill_map.get("prioritized_skills", []),
            "matched_skills": tailored_skill_map.get("matched_skills", []),
            "missing_skills": tailored_skill_map.get("missing_skills", []),
            "experience": tailored_experience,
            "projects": tailored_projects,
            "education": education,
            "certifications": certifications,
            "source_job": {
                "title": str(job.get("title", "")).strip(),
                "company": str(job.get("company", "")).strip(),
                "url": str(job.get("url", "")).strip(),
            },
        }
    except Exception as exc:
        logger.exception("Error in build_tailored_resume")
        return {
            "header": {"name": "", "email": "", "target_role": ""},
            "summary": "",
            "skills": [],
            "matched_skills": [],
            "missing_skills": [],
            "experience": [],
            "projects": [],
            "education": [],
            "certifications": [],
            "source_job": {"title": "", "company": "", "url": ""},
        }


def generate_docx(tailored_resume: Dict[str, Any], output_path: str) -> str:
    """Generate ATS-friendly DOCX resume in required section order."""
    try:
        logger.info("Generating DOCX at %s", output_path)
        document = Document()
        style = document.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)

        header = tailored_resume.get("header", {}) if isinstance(tailored_resume.get("header", {}), dict) else {}
        name = str(header.get("name", "")).strip()
        email = str(header.get("email", "")).strip()
        target_role = str(header.get("target_role", "")).strip()
        summary = str(tailored_resume.get("summary", "")).strip()

        # Header
        document.add_heading(name if name else "Candidate", level=1)
        if email or target_role:
            document.add_paragraph(" | ".join([item for item in [email, target_role] if item]))

        # Summary
        document.add_heading("Summary", level=2)
        document.add_paragraph(summary if summary else "Summary not available.")

        # Skills
        document.add_heading("Skills", level=2)
        skills = tailored_resume.get("skills", []) if isinstance(tailored_resume.get("skills", []), list) else []
        if skills:
            document.add_paragraph(", ".join([str(skill).strip() for skill in skills if str(skill).strip()]))
        else:
            document.add_paragraph("Skills not available.")

        # Experience
        document.add_heading("Experience", level=2)
        experience = tailored_resume.get("experience", []) if isinstance(tailored_resume.get("experience", []), list) else []
        if experience:
            for item in experience:
                document.add_paragraph(str(item).strip(), style="List Bullet")
        else:
            document.add_paragraph("Experience not available.")

        # Projects
        document.add_heading("Projects", level=2)
        projects = tailored_resume.get("projects", []) if isinstance(tailored_resume.get("projects", []), list) else []
        if projects:
            for item in projects:
                document.add_paragraph(str(item).strip(), style="List Bullet")
        else:
            document.add_paragraph("Projects not available.")

        # Education
        document.add_heading("Education", level=2)
        education = tailored_resume.get("education", []) if isinstance(tailored_resume.get("education", []), list) else []
        if education:
            for item in education:
                document.add_paragraph(str(item).strip(), style="List Bullet")
        else:
            document.add_paragraph("Education not available.")

        # Certifications
        document.add_heading("Certifications", level=2)
        certifications = tailored_resume.get("certifications", []) if isinstance(tailored_resume.get("certifications", []), list) else []
        if certifications:
            for item in certifications:
                document.add_paragraph(str(item).strip(), style="List Bullet")
        else:
            document.add_paragraph("Certifications not available.")

        output_dir = os.path.dirname(output_path)
        if output_dir:
            os.makedirs(output_dir, exist_ok=True)
        document.save(output_path)
        logger.info("DOCX generated successfully")
        return output_path
    except Exception as exc:
        logger.exception("Error in generate_docx")
        return ""


if __name__ == "__main__":
    try:
        logging.basicConfig(level=logging.INFO)
        logger.info("Running module self-test")
        sample_resume_data = {
            "structured_info": {
                "name": "Sample Candidate",
                "email": "sample@email.com",
                "summary": "Python developer with backend and data pipeline experience.",
                "skills": ["Python", "SQL", "AWS", "Machine Learning"],
                "experience": [
                    "Developed REST APIs using Python.",
                    "Built ETL pipelines and improved performance.",
                ],
                "projects": [
                    "Created resume matching system using NLP techniques.",
                    "Automated reporting workflows with Python scripts.",
                ],
                "education": ["B.Tech in Computer Science"],
                "certifications": ["AWS Cloud Practitioner"],
            }
        }
        sample_job = {
            "title": "Python Backend Developer",
            "company": "Example Inc",
            "description": "Need Python, SQL, and AWS skills with API development experience.",
            "url": "https://example.com/job",
        }
        tailored = build_tailored_resume(sample_resume_data, sample_job)
        output_file = os.path.join(os.path.dirname(os.path.dirname(__file__)), "tailored_resume_test.docx")
        result = generate_docx(tailored, output_file)
        logger.info("Output file: %s", result)
    except Exception as exc:
        logger.exception("Self-test failed")
